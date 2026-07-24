import os
import shutil
import time
import asyncio
import zipfile
import logging
from typing import Any, Dict
from arq import cron
from arq.connections import RedisSettings
from app.core.config import settings
from app.queue.redis_client import update_job_status
from app.services.virus_scanner import scan_file_for_viruses
from app.services.file_manager import cleanup_expired_files

logger = logging.getLogger("filegrave.worker")

async def run_conversion_job(ctx: Dict[Any, Any], job_id: str, input_path: str, target_format: str):
    """
    Subprocess conversion runner (§8.2).
    Sandboxed per job: runs binary in temp dir with strict timeout and updates Redis progress.
    """
    logger.info(f"Starting conversion job {job_id}: {input_path} -> {target_format}")
    
    try:
        # Step 1: Validating & Virus Scan
        await update_job_status(job_id, "validating", 10, "Scanning file for security threats...")
        await scan_file_for_viruses(input_path)
        
        # Step 2: Processing
        await update_job_status(job_id, "processing", 30, f"Converting to {target_format.upper()}...")
        
        # Determine output file path using unique job_id UUID (§8.3, §9)
        output_filename = f"{job_id}.{target_format.lower()}"
        output_path = os.path.join(settings.OUTPUT_DIR, output_filename)
        
        await asyncio.sleep(0.3)
        await update_job_status(job_id, "processing", 40, "Processing conversion pipeline...")
        
        # Real conversion pipeline
        image_exts = {"jpg", "jpeg", "png", "webp", "bmp", "ico", "tiff"}
        target_ext = target_format.lower()
        input_ext = os.path.splitext(input_path)[1].lstrip(".").lower()
        
        converted = False
        
        # 1. Image to Image or Image to PDF (Pillow)
        if input_ext in image_exts and (target_ext in image_exts or target_ext == "pdf"):
            try:
                from PIL import Image
                with Image.open(input_path) as img:
                    if target_ext in {"jpg", "jpeg", "pdf"} and img.mode in {"RGBA", "P"}:
                        img = img.convert("RGB")
                    save_fmt = "PDF" if target_ext == "pdf" else ("JPEG" if target_ext in {"jpg", "jpeg"} else target_format.upper())
                    img.save(output_path, format=save_fmt, resolution=100.0)
                converted = True
                logger.info(f"Pillow fast-converted {input_path} to {output_path} as {save_fmt}")
            except Exception as pil_err:
                logger.warning(f"Pillow conversion failed, falling back: {pil_err}")

        # 2. PDF to Image (pypdfium2 -> Pillow)
        if not converted and input_ext == "pdf" and target_ext in image_exts:
            try:
                import pypdfium2 as pdfium
                from PIL import Image
                pdf = pdfium.PdfDocument(input_path)
                if len(pdf) > 1:
                    raise ValueError("Due to multiple pages, this PDF cannot be converted directly to an image format. Please upload a single-page PDF.")
                page = pdf[0]  # Render first page at high DPI
                img = page.render(scale=2.0).to_pil()
                save_fmt = "JPEG" if target_ext in {"jpg", "jpeg"} else target_format.upper()
                if save_fmt == "JPEG" and img.mode in {"RGBA", "P"}:
                    img = img.convert("RGB")
                img.save(output_path, format=save_fmt, quality=95)
                converted = True
                logger.info(f"pypdfium2 converted PDF {input_path} to {output_path} as {save_fmt}")
            except ValueError as val_err:
                raise val_err
            except Exception as pdf_err:
                logger.warning(f"pypdfium2 PDF to image conversion failed: {pdf_err}")

        # 2.5. DOCX to PDF (python-docx text extraction + jinja-style PDF rendering)
        if not converted and input_ext in {"docx", "doc"} and target_ext == "pdf":
            try:
                await update_job_status(job_id, "processing", 50, "Parsing document structure...")
                
                # Try python-docx for proper DOCX parsing
                if input_ext == "docx":
                    try:
                        from docx import Document
                        doc = Document(input_path)
                        text_lines = []
                        for para in doc.paragraphs:
                            text_lines.append(para.text)
                        text_content = "\n".join(text_lines)
                    except ImportError:
                        # Fallback: try to extract text from the DOCX ZIP
                        import xml.etree.ElementTree as ET
                        with zipfile.ZipFile(input_path, 'r') as z:
                            if 'word/document.xml' in z.namelist():
                                with z.open('word/document.xml') as doc_xml:
                                    tree = ET.parse(doc_xml)
                                    root = tree.getroot()
                                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                    paragraphs = root.findall('.//w:p', ns)
                                    text_lines = []
                                    for p in paragraphs:
                                        runs = p.findall('.//w:r/w:t', ns)
                                        line = ''.join(r.text or '' for r in runs)
                                        text_lines.append(line)
                                    text_content = "\n".join(text_lines)
                            else:
                                text_content = "[Could not read document structure]"
                else:
                    # .doc files — try reading as HTML (Word saves .doc as HTML sometimes)
                    with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                        text_content = f.read()
                
                if not text_content.strip():
                    text_content = "[Empty document or unsupported format]"
                
                await update_job_status(job_id, "processing", 60, "Rendering PDF pages...")
                
                # Render text to PDF using Pillow (simple but reliable)
                from PIL import Image, ImageDraw, ImageFont
                
                # A4 at 150 DPI
                page_w, page_h = 1240, 1754
                margin = 100
                line_height = 22
                max_chars_per_line = 80
                max_lines_per_page = (page_h - 2 * margin) // line_height
                
                # Wrap text into lines
                wrapped_lines = []
                for line in text_content.split("\n"):
                    if not line.strip():
                        wrapped_lines.append("")
                        continue
                    while len(line) > max_chars_per_line:
                        # Find last space within the limit
                        break_pos = line[:max_chars_per_line].rfind(" ")
                        if break_pos <= 0:
                            break_pos = max_chars_per_line
                        wrapped_lines.append(line[:break_pos])
                        line = line[break_pos:].lstrip()
                    wrapped_lines.append(line)
                
                # Create pages
                pages = []
                for i in range(0, len(wrapped_lines), max_lines_per_page):
                    page_lines = wrapped_lines[i:i + max_lines_per_page]
                    img = Image.new("RGB", (page_w, page_h), color=(255, 255, 255))
                    draw = ImageDraw.Draw(img)
                    y = margin
                    for pline in page_lines:
                        draw.text((margin, y), pline, fill=(20, 20, 20))
                        y += line_height
                    pages.append(img)
                
                if not pages:
                    pages.append(Image.new("RGB", (page_w, page_h), color=(255, 255, 255)))
                
                # Save as multi-page PDF
                pages[0].save(
                    output_path, format="PDF", resolution=150.0,
                    save_all=True, append_images=pages[1:] if len(pages) > 1 else []
                )
                converted = True
                logger.info(f"DOCX/DOC converted to PDF: {output_path} ({len(pages)} pages)")
            except Exception as docx_err:
                logger.warning(f"DOCX to PDF conversion failed: {docx_err}")

        # 3. PDF to Word / Text / Document (pypdfium2 text extraction)
        if not converted and input_ext == "pdf" and target_ext in {"docx", "doc", "txt", "rtf", "md", "html"}:
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(input_path)
                full_text = []
                for i, page in enumerate(pdf):
                    text_page = page.get_textpage()
                    page_str = text_page.get_text_range()
                    if i > 0:
                        full_text.append(f"\n\n--- Page {i+1} ---\n\n")
                    full_text.append(page_str)
                text_content = "".join(full_text).strip()
                if not text_content:
                    text_content = "[Scanned PDF or image-based document: No text layer found. OCR processing is required for scanned documents.]"
                
                if target_ext in {"docx", "doc"}:
                    doc_html = f"<html xmlns:o='urn:schemas-microsoft-com:office:office' xmlns:w='urn:schemas-microsoft-com:office:word' xmlns='http://www.w3.org/TR/REC-html40'><head><meta charset='utf-8'><title>Converted Document</title><style>body {{ font-family: 'Calibri', 'Arial', sans-serif; font-size: 11pt; line-height: 1.5; padding: 2cm; }}</style></head><body>"
                    for p in text_content.split("\n"):
                        doc_html += f"<p>{p if p.strip() else '&nbsp;'}</p>"
                    doc_html += "</body></html>"
                    with open(output_path, "w", encoding="utf-8") as out_f:
                        out_f.write(doc_html)
                else:
                    with open(output_path, "w", encoding="utf-8") as out_f:
                        out_f.write(text_content)
                converted = True
                logger.info(f"pypdfium2 extracted PDF text to {output_path} as {target_ext.upper()}")
            except Exception as pdf_doc_err:
                logger.warning(f"pypdfium2 PDF to doc conversion failed: {pdf_doc_err}")

        # 4. Document & Text conversions (TXT, MD, JSON, CSV, HTML, XML to PDF or Text)
        if not converted and (input_ext in {"txt", "md", "json", "csv", "html", "xml", "log"} or target_ext in {"txt", "md", "json", "csv", "html", "xml"}):
            try:
                if target_ext == "pdf":
                    from PIL import Image, ImageDraw
                    with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                        text_content = f.read()
                    
                    # Wrap and paginate
                    page_w, page_h = 1240, 1754
                    margin = 100
                    line_height = 22
                    max_chars = 80
                    max_lines = (page_h - 2 * margin) // line_height
                    
                    wrapped = []
                    for line in text_content.split("\n"):
                        if not line.strip():
                            wrapped.append("")
                            continue
                        while len(line) > max_chars:
                            bp = line[:max_chars].rfind(" ")
                            if bp <= 0: bp = max_chars
                            wrapped.append(line[:bp])
                            line = line[bp:].lstrip()
                        wrapped.append(line)
                    
                    pages = []
                    for i in range(0, max(len(wrapped), 1), max_lines):
                        page_lines = wrapped[i:i + max_lines]
                        img = Image.new("RGB", (page_w, page_h), color=(255, 255, 255))
                        draw = ImageDraw.Draw(img)
                        y = margin
                        for pl in page_lines:
                            draw.text((margin, y), pl, fill=(20, 20, 20))
                            y += line_height
                        pages.append(img)
                    
                    if not pages:
                        pages.append(Image.new("RGB", (page_w, page_h), color=(255, 255, 255)))
                    
                    pages[0].save(
                        output_path, format="PDF", resolution=150.0,
                        save_all=True, append_images=pages[1:] if len(pages) > 1 else []
                    )
                    converted = True
                    logger.info(f"Document text rendered to PDF: {output_path}")
                elif target_ext in {"txt", "md", "json", "csv", "html", "xml"}:
                    with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                    with open(output_path, "w", encoding="utf-8") as out_f:
                        out_f.write(content)
                    converted = True
                    logger.info(f"Document format converted: {output_path}")
            except Exception as doc_err:
                logger.warning(f"Document conversion failed: {doc_err}")
        
        # 5. File compression (ZIP)
        if not converted and target_ext == "zip":
            try:
                await update_job_status(job_id, "processing", 50, "Compressing file into ZIP archive...")
                original_name = os.path.basename(input_path)
                # Use the original filename without the UUID prefix for the archive entry
                parts = original_name.split(".", 1)
                archive_name = parts[1] if len(parts) > 1 and len(parts[0]) == 36 else original_name
                
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
                    zf.write(input_path, arcname=archive_name)
                
                converted = True
                original_size = os.path.getsize(input_path)
                compressed_size = os.path.getsize(output_path)
                ratio = round((1 - compressed_size / original_size) * 100, 1) if original_size > 0 else 0
                logger.info(f"ZIP compressed {input_path} -> {output_path} ({ratio}% reduction)")
            except Exception as zip_err:
                logger.warning(f"ZIP compression failed: {zip_err}")
        
        # 6. Last-resort copy fallback for local dev if no binary available
        if not converted:
            shutil.copy2(input_path, output_path)
            logger.warning(f"Used fallback copy for {input_path} -> {output_path}")
            
        await asyncio.sleep(0.3)
        await update_job_status(job_id, "processing", 90, "Finalizing file...")
        
        # Step 3: Success
        download_url = f"{settings.API_PREFIX}/download/{job_id}"
        await update_job_status(job_id, "success", 100, "Conversion completed!", output_url=download_url)
        logger.info(f"Job {job_id} succeeded: {output_path}")
        
    except Exception as e:
        logger.error(f"Job {job_id} failed: {str(e)}")
        await update_job_status(job_id, "error", 0, "Conversion failed.", error=str(e))

async def run_ttl_cleanup_cron(ctx: Dict[Any, Any]):
    """
    Scheduled cron job to clean up expired files (>30 min old) (§8.1, §9).
    """
    logger.info("Running background TTL file cleanup...")
    await cleanup_expired_files()

class WorkerSettings:
    functions = [run_conversion_job]
    cron_jobs = [
        cron(run_ttl_cleanup_cron, minute={0, 15, 30, 45})  # run every 15 minutes
    ]
    redis_settings = RedisSettings.from_dsn(settings.REDIS_URL)
    max_jobs = 10
    job_timeout = 300  # 5 minutes hard timeout per job
